"""生成操盘手辅助决策方案 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

YH = 'Microsoft YaHei'  # 统一字体

def _set_cn_font(run, name=YH):
    """为 run 设置中文字体"""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

doc = Document()

# 文档属性
doc.core_properties.title = '哨响AI — 操盘手辅助决策终端 完整实施方案'
doc.core_properties.author = '哨响AI'
doc.core_properties.subject = '操盘手辅助决策方案'

style = doc.styles['Normal']
style.font.name = YH
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.3
style.element.rPr.rFonts.set(qn('w:eastAsia'), YH)
 
# 标题页
title = doc.add_heading('', level=0)
run = title.add_run('哨响AI — 操盘手辅助决策终端 完整实施方案')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x4d, 0x2e)
_set_cn_font(run)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
 
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'\n数据飞轮 · 实时终端 · 浏览器插件\n生成日期: {datetime.date.today().isoformat()}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
_set_cn_font(run)
doc.add_page_break()
 
def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.color.rgb = RGBColor(0x1a, 0x4d, 0x2e)
        _set_cn_font(r)
    return hd
 
def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.bold = bold
    _set_cn_font(r)
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p
 
def bullet(text):
    bp = doc.add_paragraph(text, style='List Bullet')
    for r in bp.runs:
        _set_cn_font(r)
    return bp
 
def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = 'Light Shading Accent 1'
    except Exception:
        pass  # 样式不存在就降级
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        t.rows[0].cells[i].text = hd
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
                _set_cn_font(r)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    _set_cn_font(r)
 
# 一、核心架构
h('一、核心架构：数据飞轮', 1)
para('本方案核心是一个自动运转的数据飞轮，持续增长训练数据，驱动操盘手决策引擎不断学习进化。由四个环节构成闭环：')
doc.add_paragraph()
 
for title_text, desc in [
    ('① 每日定时智能拉取', '每天00:05自动拉取The Odds API多庄家赔率。首次全量探测34联赛哪些有当日比赛，后续只拉活跃联赛（典型日约10-15个），节省API配额。数据落库 live_odds_raw 表作为开盘快照。'),
    ('② 比赛日实时决策', '操盘终端实时拉取最新多庄赔率，或浏览器插件推送滚球赔率。引擎调用 _live_predict 进行跨庄价值层分析，输出建仓/观望/SCAN决策卡片，含EV、凯利注码、机构意图解读。'),
    ('③ 赛果自动回填', '每6小时扫描已结束比赛，从football-data.org拉取赛果（比分+胜负），回填到 live_odds_raw 的 actual_result 字段。'),
    ('④ 训练数据增长', '每日将已回填赛果的 live_odds_raw 数据同步到 odds_features 表（30万行主表），提取开盘/收盘/drift/outcome。reverse_odds_engine 用增长后的数据再训练，操盘手持续学习。'),
]:
    p = doc.add_paragraph()
    r = p.add_run(title_text + '\n')
    r.bold = True
    r.font.size = Pt(11)
    _set_cn_font(r)
    r2 = p.add_run('   ' + desc)
    r2.font.size = Pt(10.5)
    _set_cn_font(r2)
 
doc.add_paragraph()
 
# 二、关键约束
h('二、关键约束与诚实声明', 1)
para('以下约束基于30万行严格回测验证，不可违背：')
table(['约束', '说明', '依据'], [
    ['1X2市场有效', '单庄家赔率的edge不可证伪，决策永远PASS', 'v6铁律，5方独立验证'],
    ['BET需多庄家', '必须>=2家庄的跨庄价差才能产出真实BET决策', 'bridge_service.py:1121硬编码'],
    ['配额限制', 'The Odds API免费层约390次剩余，每联赛1次请求', '实测x-requests-remaining'],
    ['不自动下单', '终端只给决策建议，用户手动操作博彩网站', '合规底线+反机器人机制'],
    ['1X2天花板', '整体预测准确率天花板52-55%', '30万行walk-forward回测'],
])
doc.add_paragraph()
 
# 三、后端数据飞轮
h('三、后端数据飞轮', 1)
h('3.1 智能定时拉取器 (pipeline/collectors/daily_collector.py)', 2)
table(['函数', '功能', '频率'], [
    ['collect_daily_odds()', '智能拉取活跃联赛多庄赔率 -> live_odds_raw', '每24h'],
    ['backfill_results()', '拉已结束比赛赛果 -> 回填actual_result', '每6h'],
    ['sync_to_odds_features()', '有赛果数据 -> odds_features增长训练集', '每24h'],
])
para('')
para('智能筛选逻辑（用户要求）：')
bullet('首次运行：全量拉取34个联赛，记录哪些有当日比赛')
bullet('缓存"活跃联赛列表"到内存')
bullet('后续定时：只拉活跃联赛（省配额），典型日约10-15个')
bullet('每次拉取后记录 x-requests-remaining，低于50时告警')
 
h('3.2 定时任务调度', 2)
para('bridge_service 无现成定时机制，用 FastAPI startup事件 + asyncio后台循环：')
bullet('daily_odds_loop：启动时立即跑一次，之后00:05定时')
bullet('result_backfill_loop：每6小时扫描回填赛果')
bullet('odds_features_sync_loop：每24小时同步训练数据')
para('幂等设计：INSERT OR REPLACE + WHERE防重复，重启安全。')
 
h('3.3 新增API接口', 2)
table(['接口', '方法', '功能'], [
    ['/api/terminal/matches', 'GET', '当天可决策比赛列表（有多庄赔率的）'],
    ['/api/terminal/analyze', 'POST', '指定比赛实时拉取多庄 -> _live_predict -> 决策卡片'],
    ['/api/terminal/ingest', 'POST', '接收插件滚球赔率 -> 实时分析 -> 决策'],
    ['/api/data-growth/stats', 'GET', '数据增长统计（行数/配额/活跃联赛）'],
])
doc.add_paragraph()
 
# 四、操盘终端UI
h('四、操盘终端UI (OperatorTerminal)', 1)
h('4.1 页面结构', 2)
bullet('顶部状态栏：API配额剩余、数据增长统计、活跃联赛数')
bullet('左栏：当天比赛列表（对阵/联赛/开赛时间/庄家数badge）')
bullet('右栏：决策卡片（赔率区+决策区+操作区）')
h('4.2 决策卡片内容', 2)
table(['区域', '内容', '数据源'], [
    ['赔率区', '多庄1X2最优价 + 隐含概率条', 'The Odds API多庄聚合'],
    ['决策标签', '建仓 / 观望 / SCAN', 'value_layer.decision'],
    ['EV/凯利', '期望价值 + 半凯利比例 + 建议注码', 'compute_value_layer + bet_core'],
    ['跨庄信号', 'consensus/disagreement/clv_beat', 'reverse_odds_engine.analyze_multi'],
    ['意图解读', '诚实防X/诱盘假X/中性 + verdict', 'classify_intent'],
    ['操作', '记录到bet_records按钮', 'betService.placeBet'],
])
h('4.3 路由与导航', 2)
bullet('router.tsx 加 /operator-terminal 路由')
bullet('Sidebar.tsx 加"操盘终端"导航项 + TerminalIcon图标')
doc.add_paragraph()
 
# 五、浏览器插件
h('五、浏览器插件 (Chrome MV3)', 1)
h('5.1 插件结构', 2)
table(['文件', '功能'], [
    ['manifest.json', 'MV3清单，权限activeTab + host_permissions'],
    ['content.js', '注入博彩网站，读取赔率DOM'],
    ['background.js', 'service worker，管理WebSocket连接'],
    ['popup.html/js', '插件弹窗：连接状态+手动推送'],
    ['adapters/generic.js', '通用1X2表格提取（fallback）'],
    ['adapters/williamhill.js', '威廉希尔DOM适配器（样板）'],
])
h('5.2 工作流', 2)
para('1. 用户打开博彩网站（B窗口） -> 插件自动识别页面')
para('2. 读1X2赔率（+让球/大小球，适配器决定）')
para('3. background.js 通过WebSocket推到 ws://localhost:9000/ws/odds_ingest')
para('4. bridge_service接收 -> _live_predict实时计算 -> broadcast到终端')
para('5. 终端实时更新决策卡片')
h('5.3 新增后端WebSocket端点', 2)
bullet('/ws/odds_ingest：接收插件推送 {home, away, source, h, d, a, score?, minute?}')
bullet('同场比赛多家推送累积 -> >=2家触发_live_predict')
bullet('结果broadcast到/ws/realtime（需ConnectionManager管理前端连接）')
doc.add_paragraph()
 
# 六、实施顺序
h('六、实施顺序（4阶段，每阶段可独立验证）', 1)
table(['阶段', '内容', '验证标准', '预估'], [
    ['阶段1 后端飞轮', 'daily_collector + 3后台循环 + API', '启动bridge->自动拉取->live_odds_raw增长', '3-4h'],
    ['阶段2 数据增长', '赛果回填 + odds_features同步', '已结束比赛回填->odds_features行数增长', '2h'],
    ['阶段3 终端UI', 'OperatorTerminal页 + terminal接口', '选比赛->看到多庄决策卡片(BET/PASS)', '4h'],
    ['阶段4 浏览器插件', 'MV3插件 + /ws/odds_ingest', '打开博彩网站->终端实时更新', '4-6h'],
])
doc.add_paragraph()
 
# 七、涉及文件
h('七、涉及文件清单', 1)
table(['文件', '操作', '说明'], [
    ['pipeline/collectors/daily_collector.py', '新建', '智能拉取+赛果回填+odds_features同步'],
    ['bridge_service.py', '修改', '3后台循环+4新API+/ws/odds_ingest+ConnectionManager'],
    ['frontend/.../OperatorTerminal/index.tsx', '新建', '操盘终端页面'],
    ['frontend/.../api.ts', '修改', '新增terminalService'],
    ['frontend/.../types/index.ts', '修改', '新增DecisionCard类型'],
    ['frontend/.../router.tsx', '修改', '加路由'],
    ['frontend/.../Sidebar.tsx', '修改', '加导航项+图标'],
    ['browser_extension/*', '新建', 'MV3插件全套'],
])
doc.add_paragraph()
 
# 八、边界
h('八、不做的事（边界）', 1)
bullet('不自动下单：插件只读赔率+推数据，终端只给决策，用户手动操作')
bullet('不推送GitHub：任务完成后留本地，推送只允许手动')
bullet('插件适配器先做generic通用提取 + 1-2家样板，不做全部庄家')
bullet('不宣称99%准确率：1X2预测物理天花板52-55%，诚实报告真实指标')
 
# 保存
output = r'C:\Users\ShXAI\Desktop\辅助决策方案.docx'
try:
    doc.save(output)
    print(f'OK 方案文档已保存: {output}')
    print(f'   文件大小: {os.path.getsize(output)} 字节')
except Exception as e:
    print(f'ERROR 保存失败: {e}')
    raise